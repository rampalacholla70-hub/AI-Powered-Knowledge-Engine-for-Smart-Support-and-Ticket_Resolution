import os
import re
import sys
import traceback
import chromadb
from pathlib import Path
from docx import Document as DocxDocument
import streamlit as st
from chromadb.utils import embedding_functions
from groq import Groq
import google.generativeai as genai
from dotenv import load_dotenv

# Load environment variables from .env file
load_dotenv()

# Get API keys from environment variables
GROQ_KEY = os.getenv("GROQ_API_KEY")
GEMINI_KEY = os.getenv("GEMINI_API_KEY")
OPENAI_KEY = os.getenv("OPENAI_API_KEY")  # used by SemanticChunker if available
os.environ["TOKENIZERS_PARALLELISM"] = "false"

# --- CONFIGURATION ---
BASE_DIR = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
DOCS_DIR = os.path.join(BASE_DIR, "docs", "dell-data")
# FIX: Use a permanent folder for the database
DB_DIR = os.path.join(BASE_DIR, "chroma_db_storage")

# Robustly get Model Name (Handling missing secrets.toml)
try:
    MODEL_NAME = st.secrets.get("EMBED_MODEL", "sentence-transformers/all-MiniLM-L6-v2")
except Exception:
    MODEL_NAME = "sentence-transformers/all-MiniLM-L6-v2"

# Chunking config (fallback defaults)
CHUNK_MAX_CHARS = int(os.getenv("CHUNK_MAX_CHARS", 800))
CHUNK_OVERLAP_CHARS = int(os.getenv("CHUNK_OVERLAP_CHARS", 200))
CHUNK_MIN_CHARS = int(os.getenv("CHUNK_MIN_CHARS", 200))

# --- API CLIENT INIT ---
if not GROQ_KEY and not GEMINI_KEY:
    print(" WARNING: No LLM API keys (Groq/Gemini) found. LLM responses will not be available.")

try:
    groq_client = Groq(api_key=GROQ_KEY) if GROQ_KEY else None
except Exception as e:
    groq_client = None
    print(f" Groq init failed: {e}")

try:
    if GEMINI_KEY:
        genai.configure(api_key=GEMINI_KEY)
        gemini_model = genai.GenerativeModel("gemini-2.5-pro")
    else:
        gemini_model = None
except Exception as e:
    gemini_model = None
    print(f" Gemini init failed: {e}")

# Try to import LangChain SemanticChunker (optional)
USE_SEMANTIC_CHUNKER = False
try:
    from langchain_experimental.text_splitter import SemanticChunker
    try:
        from langchain_huggingface import HuggingFaceEmbeddings
    except Exception:
        HuggingFaceEmbeddings = None
    SEMANTIC_CHUNKER_AVAILABLE = True
except Exception:
    SEMANTIC_CHUNKER_AVAILABLE = False


# --- HELPER CLASSES ---

class Document:
    def __init__(self, page_content, metadata=None):
        self.page_content = page_content
        self.metadata = metadata or {}

class DummyVectorDB:
    """Prevents crashes when the database is empty or failing."""
    def query(self, query_texts, n_results=4):
        return {"documents": [], "ids": [], "metadatas": []}
    def add(self, documents, metadatas, ids):
        pass


# --- TEXT PROCESSING ---

def read_docx(filepath):
    try:
        doc = DocxDocument(filepath)
        return "\n".join([p.text.strip() for p in doc.paragraphs if p.text.strip()])
    except Exception as e:
        print(f"Error reading {filepath}: {e}")
        return ""

# Fallback simple sentence splitter helper (robust, no deps)
def _split_sentences(text: str):
    text = text.strip()
    sentence_end_re = re.compile(r'(?<=[\.\?\!])\s+(?=[A-Z0-9"“])')
    parts = sentence_end_re.split(text)
    result = []
    for p in parts:
        p = p.strip()
        if not p:
            continue
        if len(p) > CHUNK_MAX_CHARS * 2:
            subparts = re.split(r',\s+', p)
            for s in subparts:
                s = s.strip()
                if s:
                    result.append(s)
        else:
            result.append(p)
    return result

# Unified chunking function: tries SemanticChunker first, else fallback
def chunk_text_to_list(text: str):
    if not text or not text.strip():
        return []

    # Try semantic chunker (requires langchain_experimental & OpenAIEmbeddings)
    if USE_SEMANTIC_CHUNKER:
        try:
            hf_embed = HuggingFaceEmbeddings(model_name=MODEL_NAME)
            splitter = SemanticChunker(hf_embed, breakpoint_threshold_type="percentile")
            docs = splitter.create_documents([text])
            out_chunks = []
            for d in docs:
                if hasattr(d, "page_content"):
                    out_chunks.append(d.page_content.strip())
                elif isinstance(d, str):
                    out_chunks.append(d.strip())
                else:
                    out_chunks.append(str(d).strip())
            if out_chunks:
                print(f"Semantic chunking successful: {len(out_chunks)} chunks")
                return out_chunks
        except Exception as e:
            print(f" SemanticChunker failed (falling back): {e}")

    # Fallback robust chunker (sentence accumulation)
    # Normalize whitespace
    text = re.sub(r'\r\n', '\n', text)
    text = re.sub(r'\n{3,}', '\n\n', text).strip()
    paragraphs = re.split(r'\n\s*\n', text)

    sentences = []
    for para in paragraphs:
        para = para.strip()
        if not para:
            continue
        lines = para.splitlines()
        for ln in lines:
            ln = ln.strip()
            if not ln:
                continue
            if re.match(r'^\#{1,6}\s+', ln) or (len(ln) <= 80 and ln.isupper()):
                sentences.append(ln)
            else:
                sentences.extend(_split_sentences(ln))

    chunks = []
    current = []
    current_len = 0
    for sent in sentences:
        s = sent.strip()
        if not s:
            continue
        s_len = len(s)
        if s_len > CHUNK_MAX_CHARS:
            parts = re.split(r'([,;])\s*', s)
            buffer = ""
            for p in parts:
                if not p:
                    continue
                if len(buffer) + len(p) + 1 > CHUNK_MAX_CHARS and buffer:
                    current.append(buffer.strip())
                    current_len += len(buffer)
                    buffer = p
                else:
                    buffer = buffer + " " + p if buffer else p
            if buffer:
                current.append(buffer.strip())
                current_len += len(buffer)
        else:
            current.append(s)
            current_len += s_len

        if current_len >= CHUNK_MAX_CHARS:
            chunk_text = " ".join(current).strip()
            chunks.append(chunk_text)
            if CHUNK_OVERLAP_CHARS > 0:
                overlap_buf = []
                overlap_len = 0
                for cs in reversed(current):
                    if overlap_len >= CHUNK_OVERLAP_CHARS:
                        break
                    overlap_buf.insert(0, cs)
                    overlap_len += len(cs)
                current = overlap_buf.copy()
                current_len = sum(len(x) for x in current)
            else:
                current = []
                current_len = 0

    if current:
        tail = " ".join(current).strip()
        if len(tail) < CHUNK_MIN_CHARS and chunks:
            chunks[-1] = chunks[-1] + "\n\n" + tail
        else:
            chunks.append(tail)

    final_chunks = [c.strip() for c in chunks if c and len(c.strip()) >= 1]
    return final_chunks

# --- VECTOR STORE FUNCTIONS ---
def create_vectorstore():
    """
    Scans DOCS_DIR and adds ONLY new files to the Chroma collection.
    Does NOT delete existing data.
    """
    global global_vectordb 

    if not os.path.exists(DOCS_DIR) or not os.listdir(DOCS_DIR):
        print(f"No .docx files found in {DOCS_DIR}")
        return DummyVectorDB()

    print(f"Updating VectorStore at: {DB_DIR}")
    embedding_func = embedding_functions.SentenceTransformerEmbeddingFunction(model_name=MODEL_NAME)
    
    # 1. Connect to DB
    chroma = chromadb.PersistentClient(path=DB_DIR)
    
    # 2. Get or Create Collection (Don't delete!)
    vectordb = chroma.get_or_create_collection(name="dell_kb", embedding_function=embedding_func)
    
    # 3. Find out what files we already have
    try:
        # Fetch only metadatas to check sources (lighter than fetching documents)
        existing_data = vectordb.get(include=['metadatas'])
        existing_sources = set()
        for m in existing_data['metadatas']:
            if m and 'source' in m:
                existing_sources.add(m['source'])
        print(f"Found {len(existing_sources)} existing documents in DB.")
    except Exception:
        existing_sources = set()

    # 4. Process only NEW files
    new_docs = []
    files_processed_count = 0
    
    for file in sorted(os.listdir(DOCS_DIR)):
        if file.lower().endswith(".docx"):
            # SKIP if already in DB
            if file in existing_sources:
                continue
            
            print(f"Processing NEW file: {file}")
            filepath = os.path.join(DOCS_DIR, file)
            text = read_docx(filepath)
            if not text:
                continue
            
            chunks = chunk_text_to_list(text)
            for idx, chunk in enumerate(chunks):
                new_docs.append(Document(page_content=chunk, metadata={"source": file, "chunk": idx}))
            
            files_processed_count += 1

    if not new_docs:
        print("No new files to add.")
        global_vectordb = vectordb
        return vectordb

    # 5. Add only the new chunks
    print(f"Adding {len(new_docs)} new chunks from {files_processed_count} files...")
    
    batch_size = 500
    # Start IDs after the current max to avoid collisions (or use UUIDs)
    # Simple strategy: Use current timestamp + index or just ensure uniqueness.
    # Safe strategy: Count existing items
    start_index = vectordb.count()
    
    for i in range(0, len(new_docs), batch_size):
        batch = new_docs[i:i+batch_size]
        vectordb.add(
            documents=[d.page_content for d in batch],
            metadatas=[d.metadata for d in batch],
            ids=[f"new_{start_index + i + j}" for j in range(len(batch))]
        )

    print(f"Successfully added {len(new_docs)} chunks.")

    # 6. Update global reference
    global_vectordb = vectordb
    
    return vectordb

def load_vectorstore():
    # FIX: Use PersistentClient
    chroma = chromadb.PersistentClient(path=DB_DIR)
    try:
        return chroma.get_collection(name="dell_kb")
    except Exception:
        print("Vectorstore not found → Creating new one...")
        return create_vectorstore()

# Global Access
global_vectordb = None

def ensure_vectordb():
    global global_vectordb
    if global_vectordb is None:
        try:
            global_vectordb = load_vectorstore()
            # Sanity check: is it empty?
            try:
                cnt = global_vectordb.count()
                if cnt == 0:
                    print("Vectorstore appears empty. Rebuilding...")
                    global_vectordb = create_vectorstore()
            except Exception:
                pass
        except Exception as e:
            print("Error loading vectorstore:", e)
            traceback.print_exc()
            global_vectordb = DummyVectorDB()
    return global_vectordb

# --- RAG EXECUTION ---

def get_ai_response(query, docs):
    print(f"DEBUG: Generating AI response for: {query}") 
    context = "\n\n".join(docs) if docs else ""
    prompt = f"""You are a Dell technical support assistant.
    Use the following Dell knowledge base context to answer the user question.
    
    Context:
    {context}
    
    User Question: {query}
    
    Answer concisely and helpfully."""

    # Prefer Groq if available
    if groq_client:
        try:
            print("DEBUG: Calling Groq...")
            response = groq_client.chat.completions.create(
                model="llama-3.3-70b-versatile",
                messages=[{"role": "user", "content": prompt}],
                temperature=0.3
            )
            ans = response.choices[0].message.content.strip()
            print("DEBUG: Groq Success!")
            return ans
        except Exception as e:
            print(f"DEBUG: Groq failed: {e}")

    # Next: Gemini
    if gemini_model:
        try:
            print("DEBUG: Calling Gemini...")
            resp = gemini_model.generate_content(prompt)
            # Gemini returns resp.text in many wrappers
            ans = getattr(resp, "text", str(resp)).strip()
            print("DEBUG: Gemini Success!")
            return ans
        except Exception as e:
            print(f"DEBUG: Gemini failed: {e}")

    return "No AI available."

def run_rag(query: str, n_results: int = 4):
    """
    Main callable used by frontend. Always returns dict:
      { "user_query": str, "rag_response": str, "doc_context": list }
    """
    if not query or not isinstance(query, str):
        return {"user_query": query, "rag_response": "Invalid Query", "doc_context": []}

    vectordb = ensure_vectordb()
    
    # Safety check for dummy DB
    if isinstance(vectordb, DummyVectorDB):
        return {"user_query": query, "rag_response": "Knowledge base is empty. Upload documents first.", "doc_context": []}

    try:
        results = vectordb.query(query_texts=[query], n_results=n_results)
        # Chroma returns documents as a list-of-lists or similar; handle safely
        docs = []
        if results and results.get("documents"):
            docs = results["documents"][0]
        elif results and results.get("matches"):
            # fallback in some clients
            docs = [m['document'] for m in results['matches'][:n_results]]
        docs = docs or []

        if not docs:
            ai_response = " I couldn't find any relevant information in the KB for your query."
        else:
            ai_response = get_ai_response(query, docs)
        
        return {
            "user_query": query,
            "rag_response": ai_response,
            "doc_context": docs
        }
    except Exception as e:
        print(f"RAG Error: {e}")
        return {"user_query": query, "rag_response": "Error processing request.", "doc_context": []}

if __name__ == "__main__":
    print("Test run . This will build or load the vectorstore and run a sample query.")
    vect = ensure_vectordb()
    sample = run_rag("My laptop won't turn on", n_results=2)
    print(sample)
