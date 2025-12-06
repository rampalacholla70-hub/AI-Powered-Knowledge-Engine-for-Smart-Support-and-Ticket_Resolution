from pathlib import Path
import os
import re
from docx import Document as DocxDocument
from chromadb import Client
from dotenv import load_dotenv

#  embeddings + semantic chunking
from langchain_experimental.text_splitter import SemanticChunker
#from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_huggingface import HuggingFaceEmbeddings


# Load environment
BASE_DIR = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
ENV_PATH = os.path.join(BASE_DIR, ".env")

print("Loading .env from:", ENV_PATH)
load_dotenv(ENV_PATH)

DOCS_DIR = os.path.join(BASE_DIR, "docs", "dell-data")
os.environ["TOKENIZERS_PARALLELISM"] = "false"



# Document structure

class Document:
    def __init__(self, page_content, metadata=None):
        self.page_content = page_content
        self.metadata = metadata or {}



# Read .docx

def read_docx(filepath):
    doc = DocxDocument(filepath)
    return "\n".join([p.text.strip() for p in doc.paragraphs if p.text.strip()])



# Semantic Chunking (HuggingFace)

def semantic_chunk_text(text: str):
    if not text.strip():
        return []

    # Same model used for both chunking & vectorstore
    embed_model = HuggingFaceEmbeddings(
        model_name="sentence-transformers/all-MiniLM-L6-v2"
    )

    # Semantic chunker  byE HF embeddings
    splitter = SemanticChunker(
        embed_model,
        breakpoint_threshold_type="percentile"
    )

    docs = splitter.create_documents([text])

    return [d.page_content.strip() for d in docs]



# Vectorstore Creation (Chroma + HuggingFace)

def create_vectorstore():

    if not os.path.exists(DOCS_DIR) or not os.listdir(DOCS_DIR):
        raise FileNotFoundError(f"No DOCX files in {DOCS_DIR}")

    docs = []

    for file in sorted(os.listdir(DOCS_DIR)):
        if not file.lower().endswith(".docx"):
            continue

        filepath = os.path.join(DOCS_DIR, file)
        print(f"\nProcessing file: {file}")

        text = read_docx(filepath)
        if not text:
            print(f"[WARN] Empty file skipped: {file}")
            continue

        chunks = semantic_chunk_text(text)

        for idx, chunk in enumerate(chunks):
            docs.append(Document(
                page_content=chunk,
                metadata={"source": file, "chunk": idx}
            ))

        print(f" → {len(chunks)} semantic chunks")

    print(f"\nTotal semantic chunks: {len(docs)}")

    # SAME embedding model used across entire pipeline
    emb_model = HuggingFaceEmbeddings(
        model_name="sentence-transformers/all-MiniLM-L6-v2"
    )

    chroma = Client()

    # Remove old store to rebuild fresh
    try:
        chroma.delete_collection("dell_kb")
    except:
        pass

    vectordb = chroma.create_collection(
        name="dell_kb",
        embedding_function=emb_model.embed_documents  # IMPORTANT FIX
    )

    vectordb.add(
        documents=[d.page_content for d in docs],
        metadatas=[d.metadata for d in docs],
        ids=[str(i) for i in range(len(docs))]
    )

    print("\n[✓] Vectorstore created successfully.\n")
    return vectordb



# Manual run

if __name__ == "__main__":
    print("Rebuilding vectorstore...")
    create_vectorstore()
    print("Rebuild complete.")
